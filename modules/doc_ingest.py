#!/usr/bin/env python3
"""Document ingestion for the local knowledge base.

The ingester scans:
- knowledge_base/raw_docs/ for source documents copied or extracted from archives
- knowledge_base/domain_docs/ for hand-written Markdown knowledge

It extracts text from PDF/DOCX/TXT/MD, writes normalized Markdown into
knowledge_base/processed_docs/, and rebuilds docs_manifest.json plus chunks.jsonl.
No external API is called.
"""

from __future__ import annotations

import hashlib
import json
import re
from datetime import datetime
from pathlib import Path
from typing import Callable

ROOT = Path(__file__).resolve().parent.parent
RAW_DIR = ROOT / "knowledge_base" / "raw_docs"
DOMAIN_DIR = ROOT / "knowledge_base" / "domain_docs"
PROCESSED_DIR = ROOT / "knowledge_base" / "processed_docs"
INDEX_DIR = ROOT / "knowledge_base" / "index"

CHUNK_MIN = 800
CHUNK_MAX = 1200
OVERLAP = 120

SUPPORTED_SUFFIXES = {".pdf", ".docx", ".txt", ".md"}
RAW_SKIP_PARTS = {"archives", "other", "spreadsheets", "images", "__pycache__"}


def _ts() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def _clean_text(text: str) -> str:
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _safe_stem(path: Path) -> str:
    stem = re.sub(r"[^\w\-.\u4e00-\u9fff]+", "_", path.stem, flags=re.UNICODE)
    return stem.strip("._") or "document"


def _rel(path: Path) -> str:
    return str(path.resolve().relative_to(ROOT.resolve())).replace("\\", "/")


def _make_doc_id(source_path: Path) -> str:
    digest = hashlib.sha1(_rel(source_path).encode("utf-8")).hexdigest()[:12]
    return f"DOC-{digest}"


def _read_text_fallback(path: Path) -> str:
    for enc in ("utf-8-sig", "utf-8", "gb18030", "gbk", "big5"):
        try:
            return path.read_text(encoding=enc)
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="replace")


def _load_source_metadata() -> dict[str, dict]:
    metadata_path = INDEX_DIR / "source_metadata.jsonl"
    if not metadata_path.exists():
        return {}
    records: dict[str, dict] = {}
    with metadata_path.open("r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if not line:
                continue
            try:
                row = json.loads(line)
            except json.JSONDecodeError:
                continue
            rel_path = row.get("source_file")
            if rel_path:
                records[rel_path] = row
    return records


def extract_pdf(filepath: Path) -> tuple[str, int]:
    import fitz  # PyMuPDF

    doc = fitz.open(str(filepath))
    pages: list[str] = []
    try:
        for i, page in enumerate(doc, 1):
            text = _clean_text(page.get_text(sort=True))
            pages.append(f"## Page {i}\n\n{text or '(No extractable text on this page.)'}")
    finally:
        doc.close()
    return "\n\n".join(pages), len(pages)


def extract_docx(filepath: Path) -> tuple[str, int]:
    from docx import Document

    doc = Document(str(filepath))
    blocks: list[str] = []
    section_count = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name if para.style else ""
        if style.startswith("Heading"):
            match = re.search(r"\d+", style)
            level = min(int(match.group()) if match else 1, 6)
            blocks.append(f"{'#' * level} {text}")
            section_count += 1
        else:
            blocks.append(text)

    for table in doc.tables:
        rows: list[list[str]] = []
        for row in table.rows:
            rows.append([cell.text.strip().replace("\n", " ") for cell in row.cells])
        if rows:
            width = max(len(r) for r in rows)
            normalized = [r + [""] * (width - len(r)) for r in rows]
            lines = ["| " + " | ".join(normalized[0]) + " |"]
            lines.append("| " + " | ".join(["---"] * width) + " |")
            for r in normalized[1:]:
                lines.append("| " + " | ".join(r) + " |")
            blocks.append("\n".join(lines))
            section_count += 1

    return _clean_text("\n\n".join(blocks)), max(section_count, 1)


def extract_text(filepath: Path) -> tuple[str, int]:
    text = _clean_text(_read_text_fallback(filepath))
    section_count = len(re.findall(r"^#{1,6}\s", text, re.MULTILINE))
    return text, max(section_count, 1)


EXTRACTORS: dict[str, Callable[[Path], tuple[str, int]]] = {
    ".pdf": extract_pdf,
    ".docx": extract_docx,
    ".txt": extract_text,
    ".md": extract_text,
}

OUTPUT_SUBDIR = {
    ".pdf": "pdf",
    ".docx": "word",
    ".txt": "text",
    ".md": "text",
}


def chunk_text(text: str, chunk_min: int = CHUNK_MIN,
               chunk_max: int = CHUNK_MAX, overlap: int = OVERLAP) -> list[str]:
    if len(text) <= chunk_max:
        return [text]

    chunks: list[str] = []
    buf = ""
    paragraphs = re.split(r"\n\n+", text)

    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        if len(para) > chunk_max:
            if buf.strip():
                chunks.append(buf.strip())
                buf = ""
            step = max(chunk_max - overlap, 1)
            for i in range(0, len(para), step):
                piece = para[i:i + chunk_max].strip()
                if piece:
                    chunks.append(piece)
            continue
        if buf and len(buf) + len(para) + 2 > chunk_max:
            chunks.append(buf.strip())
            buf = buf[-overlap:] if len(buf) > overlap else ""
        buf += para + "\n\n"

    if buf.strip():
        chunks.append(buf.strip())

    merged: list[str] = []
    for c in chunks:
        if len(c) < chunk_min and merged:
            merged[-1] += "\n\n" + c
        else:
            merged.append(c)
    return merged or [text]


def _raw_files() -> list[Path]:
    files: list[Path] = []
    if not RAW_DIR.exists():
        return files
    for f in RAW_DIR.rglob("*"):
        if not f.is_file() or f.suffix.lower() not in SUPPORTED_SUFFIXES:
            continue
        rel_parts = set(f.relative_to(RAW_DIR).parts[:-1])
        if rel_parts & RAW_SKIP_PARTS:
            continue
        files.append(f)
    return files


def _domain_files() -> list[Path]:
    if not DOMAIN_DIR.exists():
        return []
    return [
        f for f in DOMAIN_DIR.rglob("*.md")
        if f.is_file() and f.name.lower() != "readme.md" and "index" not in f.parts
    ]


def scan_documents() -> list[dict]:
    docs: list[dict] = []
    for f in _raw_files():
        docs.append({"source_path": f, "suffix": f.suffix.lower(), "category": "raw"})
    for f in _domain_files():
        docs.append({"source_path": f, "suffix": ".md", "category": "domain"})
    docs.sort(key=lambda d: _rel(d["source_path"]))
    return docs


def _write_processed(source: Path, suffix: str, text: str) -> Path:
    out_dir = PROCESSED_DIR / OUTPUT_SUBDIR.get(suffix, "text")
    out_dir.mkdir(parents=True, exist_ok=True)
    digest = hashlib.sha1(_rel(source).encode("utf-8")).hexdigest()[:8]
    out_path = out_dir / f"{_safe_stem(source)}_{digest}.md"
    out_path.write_text(text, encoding="utf-8")
    return out_path


def process_document(doc_info: dict, metadata_by_source: dict[str, dict] | None = None) -> dict | None:
    metadata_by_source = metadata_by_source or {}
    source: Path = doc_info["source_path"]
    suffix = doc_info["suffix"]
    category = doc_info["category"]

    extractor = EXTRACTORS.get(suffix)
    if not extractor:
        print(f"  [skip] Unsupported file type: {source.name}")
        return None

    print(f"  [ingest] {source.name} ({suffix.lstrip('.')})")
    try:
        text, count = extractor(source)
    except Exception as exc:
        print(f"  [error] Extract failed: {source.name} -> {exc}")
        return None
    if not text.strip():
        print(f"  [skip] No extractable text: {source.name}")
        return None

    out_path = _write_processed(source, suffix, text)
    doc_id = _make_doc_id(source)
    source_rel = _rel(source)
    source_meta = metadata_by_source.get(source_rel, {})

    manifest_entry = {
        "doc_id": doc_id,
        "source_file": source_rel,
        "processed_file": _rel(out_path),
        "file_type": suffix.lstrip("."),
        "file_name": source.name,
        "category": category,
        "char_count": len(text),
        "processed_time": _ts(),
        "source_reliability": source_meta.get("source_reliability", "unreviewed"),
        "safety_class": source_meta.get("safety_class", "general_reference"),
        "applicable_task": source_meta.get("applicable_task", ["search", "high_level_reference"]),
        "not_applicable_task": source_meta.get("not_applicable_task", []),
    }
    if suffix == ".pdf":
        manifest_entry["page_count"] = count
    else:
        manifest_entry["section_count"] = count

    chunks = []
    for i, chunk in enumerate(chunk_text(text)):
        chunks.append({
            "chunk_id": f"{doc_id}-C{i:04d}",
            "doc_id": doc_id,
            "source_file": source_rel,
            "processed_file": manifest_entry["processed_file"],
            "chunk_index": i,
            "page": -1,
            "text": chunk,
            "char_count": len(chunk),
            "source_reliability": manifest_entry["source_reliability"],
            "safety_class": manifest_entry["safety_class"],
            "applicable_task": manifest_entry["applicable_task"],
            "not_applicable_task": manifest_entry["not_applicable_task"],
        })

    print(f"       -> {out_path.name} ({len(text)} chars, {len(chunks)} chunks)")
    return {"manifest": manifest_entry, "chunks": chunks}


def ingest_all() -> dict:
    INDEX_DIR.mkdir(parents=True, exist_ok=True)
    for sub in ("pdf", "word", "text"):
        (PROCESSED_DIR / sub).mkdir(parents=True, exist_ok=True)

    print("=" * 60)
    print("  Knowledge base ingestion")
    print(f"  Time: {_ts()}")
    print("=" * 60)

    docs = scan_documents()
    print(f"  Found {len(docs)} supported documents")
    if not docs:
        return {"manifest": [], "chunks": []}

    metadata_by_source = _load_source_metadata()
    manifest_entries: list[dict] = []
    all_chunks: list[dict] = []
    for doc in docs:
        result = process_document(doc, metadata_by_source)
        if result:
            manifest_entries.append(result["manifest"])
            all_chunks.extend(result["chunks"])

    manifest_data = {
        "generated_at": _ts(),
        "total_docs": len(manifest_entries),
        "total_chunks": len(all_chunks),
        "docs": manifest_entries,
    }
    (INDEX_DIR / "docs_manifest.json").write_text(
        json.dumps(manifest_data, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    with (INDEX_DIR / "chunks.jsonl").open("w", encoding="utf-8") as f:
        for chunk in all_chunks:
            f.write(json.dumps(chunk, ensure_ascii=False) + "\n")

    print("\n  Done")
    print(f"  Documents: {len(manifest_entries)}")
    print(f"  Chunks:    {len(all_chunks)}")
    return {"manifest": manifest_entries, "chunks": all_chunks}


def list_manifest() -> list[dict]:
    path = INDEX_DIR / "docs_manifest.json"
    if not path.exists():
        return []
    return json.loads(path.read_text(encoding="utf-8")).get("docs", [])


def show_manifest() -> None:
    docs = list_manifest()
    if not docs:
        print("No indexed documents. Run: python scripts/ingest_documents.py")
        return
    print(f"\nIndexed documents: {len(docs)}\n")
    for i, d in enumerate(docs, 1):
        count_key = "page_count" if "page_count" in d else "section_count"
        print(f"{i:>3}. [{d.get('file_type', '?').upper()}] {d.get('file_name', '?')}")
        print(f"     {d.get('char_count', 0)} chars | {count_key}: {d.get(count_key, '-')}")
        print(f"     {d.get('source_file', '')}")


if __name__ == "__main__":
    ingest_all()
